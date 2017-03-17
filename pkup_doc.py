#! /usr/bin/env python
# -*- coding: utf-8 -*-

# this is a script generating *.docx Word document with PKUP report

from __future__ import unicode_literals
from docx import Document
from docx.shared import Inches
from datetime import date, datetime, timedelta
import sys, getopt

def generate_report():
    create_document()
    set_file_path()
    generate_heading()
    generate_dates()
    generate_personal_data()
    generate_work_results()
    generate_legal_text()
    save_report()

def create_document():
    global document
    document = Document()

def set_file_path():
    global dir_path
    dir_path = './'

def generate_heading():
    document.add_heading('Raport Miesięczny Pracy Twórczej', 0)

def generate_dates():
    generate_bold_text('Okres')
    table = document.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    date_twenty_days_ago = datetime.today() - timedelta(days=20)
    from_date = "19.%d.%d" % (date_twenty_days_ago.month, date_twenty_days_ago.year)
    to_date = date.today().strftime('%d.%m.%Y')
    date_range = "%s - %s" % (from_date, to_date)
    set_row(table, 0, "Okres raportowany od 19 do 18 dnia miesiąca", date_range)
    set_row(table, 1, "Data sporządzenia raportu", to_date)

def generate_personal_data():
    generate_bold_text('Dane osobowe')
    table = document.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    set_row(table, 0, "Imię i nazwisko pracownika", employee_name)
    set_row(table, 1, "Stanowisko służbowe", employee_role)
    set_row(table, 2, "Dział", employee_dept)
    set_row(table, 3, "Imię i nazwisko przełożonego pracownika", employee_mgr)

def generate_work_results():
    generate_bold_text('Raportowane wyniki pracy twórczej')
    table = document.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    set_row(table, 0, "Rodzaj wyniku pracy twórczej", "Opis wykonywanych czynności")
    set_row(table, 1, "Prace programistyczne i specyfikacja techniczna", report_text)

def set_row(table, row_number, label, value):
    table.rows[row_number].cells[0].text = label
    table.rows[row_number].cells[1].text = value

def generate_legal_text():
    text = "Załącznik nr. 1 do Regulaminu Działalności Twórczej z dnia 28.12.2007"
    generate_bold_text(text)

def generate_bold_text(text):
    paragraph = document.add_paragraph()
    paragraph.add_run('\n' + text).bold = True

def save_report():
    name, surname = employee_name.split(" ")
    doc_prefix = (name[:1] + surname[:2]).upper()
    file_name = "%s_raport_pkup_%d_%d.docx" % (doc_prefix, date.today().month, date.today().year)
    full_file_path = dir_path + file_name
    document.save(full_file_path)
    print('report saved successfully to: %s' % (full_file_path))

def main(argv):
    global employee_name
    global employee_role
    global employee_dept
    global employee_mgr
    global report_text

    help_message = 'pkup_doc.py -n <name and surname> -r <role> -d <department> -m <manager> -t <text>'

    try:
        opts, args = getopt.getopt(argv,"n:r:d:m:t",["name=","role=","department=","manager=","text="])
    except getopt.GetoptError:
        print(help_message)
        sys.exit(2)
    for opt, arg in opts:
        if opt == '-h':
            print(help_message)
            sys.exit()
        elif opt in ("-n", "--name"):
            employee_name = arg.decode('utf-8')
        elif opt in ("-r", "--role"):
            employee_role = arg.decode('utf-8')
        elif opt in ("-d", "--department"):
            employee_dept = arg.decode('utf-8')
        elif opt in ("-m", "--manager"):
            employee_mgr = arg.decode('utf-8')
        elif opt in ("-t", "--text"):
            report_text = arg.decode('utf-8')

    generate_report()

if __name__ == "__main__":
    main(sys.argv[1:])
